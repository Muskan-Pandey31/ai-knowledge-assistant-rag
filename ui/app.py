import streamlit as st
import io
import pandas as pd
from docx import Document

from backend.rag_pipeline import VectorStore
from backend.llm_handler import generate_answer


# 🔹 Chunking with overlap
def chunk_text(text, chunk_size=100, overlap=20):
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)
    
    return chunks


# 🔹 Section-based chunking
def process_text(text):
    sections = text.split("\n\n")
    chunks = []
    
    for section in sections:
        chunks.extend(chunk_text(section))
    
    return chunks


# 🔹 Caching file processing
@st.cache_data
def process_file(file, file_type):
    if file_type == "txt":
        text = file.read().decode("utf-8")
        return process_text(text)

    elif file_type == "csv":
        df = pd.read_csv(file)
        return df.astype(str).apply(lambda row: " ".join(row), axis=1).tolist()

    return []


# Initialize store
if "store" not in st.session_state:
    st.session_state.store = VectorStore()

store = st.session_state.store

st.title("AI Knowledge Assistant")

# 📂 File Upload
uploaded_file = st.file_uploader("Upload a TXT or CSV file")

if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1]

    chunks = process_file(uploaded_file, file_type)

    if chunks:
        store.add_documents(chunks)
        st.success("File processed successfully!")
    else:
        st.error("Only TXT and CSV supported for now")


# ❓ Ask Question
query = st.text_input("Ask a question")

answer = ""

if query:
    results = store.search(query)

    if not results:
        answer = "No relevant data found"
    else:
        context = " ".join(results)
        answer = generate_answer(context, query)

    st.write("### Answer:")
    st.write(answer)


# 📥 Download Functions
def create_word(answer):
    doc = Document()
    doc.add_heading("AI Answer", 0)
    doc.add_paragraph(answer)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_excel(answer):
    df = pd.DataFrame({"Answer": [answer]})
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False)
    buffer.seek(0)
    return buffer


# ⬇️ Download Buttons
if answer:
    st.download_button(
        "Download as Word",
        create_word(answer),
        file_name="answer.docx"
    )

    st.download_button(
        "Download as Excel",
        create_excel(answer),
        file_name="answer.xlsx"
    )